"""Résumé file import/export — parse an uploaded résumé to plain text, and
export the tailored text to .txt / .docx.

Uses python-docx / pypdf when available for fidelity, with a pure-stdlib
fallback for .txt and .docx (zipfile). PDF requires pypdf. These are OPTIONAL —
the security core (core.py) stays pure stdlib; this is a convenience layer.
The uploaded file is the user's OWN résumé (source of truth), not untrusted
external data, so it is parsed directly (with a size cap as a safety valve).
"""

import io
import re
import zipfile
import xml.sax.saxutils as _sx

MAX_UPLOAD_BYTES = 8 * 1024 * 1024  # 8 MB


class ParseError(Exception):
    pass


# --- import (file -> text) --------------------------------------------------

def parse_resume(filename, data):
    """Return plain text extracted from an uploaded résumé file (bytes)."""
    if data is None:
        raise ParseError("Empty file.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise ParseError("File too large (max 8 MB).")
    name = (filename or "").lower()
    if name.endswith(".docx"):
        return _docx_to_text(data)
    if name.endswith(".pdf"):
        return _pdf_to_text(data)
    # .txt / .md / unknown -> best-effort text decode
    return _decode_text(data)


def _decode_text(data):
    for enc in ("utf-8-sig", "utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc).strip()
        except Exception:
            continue
    return data.decode("utf-8", "replace").strip()


def _docx_to_text(data):
    # Prefer python-docx; fall back to stdlib zipfile parsing of word/document.xml.
    try:
        import docx  # python-docx
        doc = docx.Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs).strip()
    except ImportError:
        pass
    except Exception as exc:
        raise ParseError(f"Could not read .docx: {exc}")
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            xml = z.read("word/document.xml").decode("utf-8", "replace")
    except Exception as exc:
        raise ParseError(f"Not a valid .docx file: {exc}")
    lines = []
    for para in re.split(r"</w:p>", xml):
        runs = re.findall(r"<w:t[^>]*>(.*?)</w:t>", para, re.S)
        lines.append("".join(_sx.unescape(r) for r in runs))
    return "\n".join(lines).strip()


def _pdf_to_text(data):
    try:
        import pypdf
    except ImportError:
        raise ParseError(
            "PDF parsing needs pypdf (pip install pypdf). "
            "Or paste the text, or upload a .docx / .txt."
        )
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    except Exception as exc:
        raise ParseError(f"Could not read PDF: {exc}")


# --- export (text -> file) --------------------------------------------------

def to_txt(text):
    return (text or "").encode("utf-8")


def to_docx(text):
    """Return .docx bytes for the given text (one paragraph per line)."""
    try:
        import docx
        doc = docx.Document()
        for line in (text or "").split("\n"):
            doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return _minimal_docx(text)


def _minimal_docx(text):
    """Build a minimal but valid .docx with only stdlib (zipfile)."""
    paras = "".join(
        '<w:p><w:r><w:t xml:space="preserve">{}</w:t></w:r></w:p>'.format(_sx.escape(line))
        for line in (text or "").split("\n")
    )
    document = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:body>' + paras + '<w:sectPr/></w:body></w:document>'
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '</Types>'
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/></Relationships>'
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", rels)
        z.writestr("word/document.xml", document)
    return buf.getvalue()
